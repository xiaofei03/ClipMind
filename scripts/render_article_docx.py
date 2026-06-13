from __future__ import annotations

import json
import math
import re
import subprocess
import sys
import urllib.request
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


SESSION_DIR = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else None
if not SESSION_DIR:
    print("Usage: python render_article_docx.py <session-dir>", file=sys.stderr)
    sys.exit(2)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(read_text(path))


def clean_public_text(value: str) -> str:
    text = str(value or "")
    text = re.sub(r"证据图号[:：][^。！？\n]*(?:[。！？]|$)", "", text)
    text = re.sub(r"(截图证据|证据截图|证据帧|证据图|证据)", "关键画面", text)
    text = re.sub(r"(这个视频|本视频|该视频|视频中|博主|作者|UP主|up主)", "", text)
    text = re.sub(r"图\d+验证时需[^。！？\n]*(?:[。！？]|$)", "", text)
    text = re.sub(r"因为浏览器缓存可能导致旧版\s*JS\s*生效[^。！？\n]*(?:[。！？]|$)", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def strip_front_matter(markdown: str) -> str:
    return re.sub(r"^---[\s\S]*?---\s*", "", markdown or "").strip()


def strip_evidence_map(markdown: str) -> str:
    text = strip_front_matter(markdown)
    marker = "\n---\n\n"
    idx = text.find(marker)
    if idx >= 0:
        text = text[idx + len(marker) :].strip()
    text = re.sub(r"\n---\n\n## 证据文件[\s\S]*$", "", text).strip()
    text = re.sub(r"^##\s*关键截图地图[\s\S]*?(?=^##\s+|\n---\n|\s*$)", "", text, flags=re.M).strip()
    return text


def load_title(markdown: str, context: dict, metadata: dict) -> str:
    match = re.search(r"^#\s+(.+)$", strip_evidence_map(markdown), flags=re.M)
    markdown_title = clean_public_text(match.group(1).strip()) if match else "学习图文笔记"
    raw = (
        context.get("metadata", {}).get("title")
        or context.get("metadata", {}).get("fulltitle")
        or context.get("bilibiliApi", {}).get("title")
        or context.get("douyinApi", {}).get("title")
        or metadata.get("title")
        or metadata.get("fulltitle")
        or markdown_title
    )
    title = clean_public_text(raw)
    title = re.sub(r"^【[^】]{1,24}】\s*", "", title)
    title = re.sub(r"^《(.{1,40})》\s*", r"\1 ", title)
    title = re.sub(r"[|｜].*$", "", title).strip()
    return title or "学习图文笔记"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.5, first_indent=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_indent:
        pf.first_line_indent = Cm(0.74)


def add_text_run(paragraph, text: str, bold=False, code=False):
    parts = re.split(r"(`[^`]+`|\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, east_asia="Consolas", latin="Consolas", size=10, color=(17, 24, 39))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=10.5, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=10.5, bold=bold)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    set_paragraph_spacing(paragraph, before=16 if level <= 2 else 10, after=8, line=1.25)
    run = paragraph.add_run(clean_public_text(text))
    if level == 1:
        set_run_font(run, east_asia="黑体", latin="Arial", size=18, bold=True, color=(32, 36, 43))
    elif level == 2:
        set_run_font(run, east_asia="黑体", latin="Arial", size=15, bold=True, color=(32, 36, 43))
    else:
        set_run_font(run, east_asia="黑体", latin="Arial", size=13, bold=True, color=(32, 36, 43))


def image_size_cm(path: Path, max_width_cm: float, max_height_cm: float | None = None) -> tuple[float, float]:
    with Image.open(path) as img:
        w, h = img.size
    ratio = h / w if w else 0.62
    width = max_width_cm
    height = width * ratio
    if max_height_cm and height > max_height_cm:
        height = max_height_cm
        width = height / ratio if ratio else max_width_cm
    return width, height


def add_picture(document: Document, image_path: Path, caption: str | None = None, width_cm=14.8, max_height_cm=10.5):
    if not image_path.exists():
        return
    document.add_paragraph().paragraph_format.keep_with_next = True
    width, _height = image_size_cm(image_path, width_cm, max_height_cm)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = bool(caption)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width))
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, after=8, line=1.2)
        r = cap.add_run(caption)
        set_run_font(r, east_asia="宋体", latin="Times New Roman", size=9.5, bold=True, color=(180, 35, 24))


def download_cover(url: str, out_path: Path) -> Path | None:
    if not url or not re.match(r"^https?://", url):
        return None
    if out_path.exists() and out_path.stat().st_size > 1024:
        return out_path
    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125 Safari/537.36",
            "Referer": "https://www.bilibili.com/",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            out_path.write_bytes(resp.read())
        return out_path if out_path.exists() and out_path.stat().st_size > 1024 else None
    except Exception:
        return None


def extract_video_cover(video_path: str | None, out_path: Path) -> Path | None:
    if not video_path or not Path(video_path).exists():
        return None
    if out_path.exists() and out_path.stat().st_size > 1024:
        return out_path
    result = subprocess.run(
        ["ffmpeg", "-y", "-ss", "0.5", "-i", video_path, "-frames:v", "1", "-q:v", "2", str(out_path)],
        cwd=str(SESSION_DIR),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="ignore",
    )
    return out_path if result.returncode == 0 and out_path.exists() else None


def resolve_cover(context: dict, metadata: dict) -> Path | None:
    candidates = [
        context.get("metadata", {}).get("thumbnail"),
        context.get("bilibiliApi", {}).get("coverUrl"),
        context.get("douyinApi", {}).get("coverUrl"),
        metadata.get("thumbnail"),
    ]
    thumbs = metadata.get("thumbnails") or []
    for item in thumbs:
        if isinstance(item, dict) and item.get("url"):
            candidates.append(item["url"])
            break
    for url in candidates:
        ext = ".png" if re.search(r"\.png(?:[?#]|$)", str(url or ""), re.I) else ".jpg"
        downloaded = download_cover(str(url or ""), SESSION_DIR / f"article_thumbnail_docx{ext}")
        if downloaded:
            return downloaded
    return extract_video_cover(context.get("videoPath"), SESSION_DIR / "article_cover_docx.jpg")


def sorted_evidence(evidence_json: dict) -> list[dict]:
    items = evidence_json.get("selectedEvidence") or []
    result = []
    for idx, item in enumerate(items, start=1):
        if item.get("assetPath") and Path(item["assetPath"]).exists():
            item = dict(item)
            item["_displayFigure"] = f"图{idx}"
            result.append(item)
    return result


def categories(item: dict) -> set[str]:
    return set(item.get("categories") or [])


def bucket_for_heading(heading: str) -> str:
    text = heading or ""
    if re.search(r"工具|平台|场景", text):
        return "tools"
    if re.search(r"步骤\s*1|需求|Prompt|提示词", text, re.I):
        return "step1"
    if re.search(r"步骤\s*2|骨架|基础|生成|预览", text):
        return "step2"
    if re.search(r"步骤\s*3|样式|设计|组件|动效|背景|结果", text):
        return "step3"
    if re.search(r"Prompt|提示词|模板", text, re.I):
        return "prompts"
    if re.search(r"方法|复盘|经验", text):
        return "methods"
    if re.search(r"避坑|风险|错误|注意", text):
        return "pitfalls"
    return ""


def pick_evidence_for_bucket(evidence: list[dict], bucket: str) -> list[dict]:
    if bucket == "tools":
        wanted = [x for x in evidence if "tool" in categories(x)]
    elif bucket == "step1":
        wanted = [x for x in evidence if categories(x) & {"prompt", "workflow"}]
    elif bucket == "step2":
        wanted = [x for x in evidence if categories(x) & {"code", "workflow"}]
    elif bucket == "step3":
        wanted = [x for x in evidence if categories(x) & {"design", "component", "result"}]
    elif bucket == "prompts":
        wanted = [x for x in evidence if "prompt" in categories(x)]
    elif bucket == "methods":
        wanted = [x for x in evidence if categories(x) & {"design", "workflow"}]
    elif bucket == "pitfalls":
        wanted = [x for x in evidence if categories(x) & {"problem", "code"}]
    else:
        wanted = []
    return wanted[:2]


def add_evidence_block(document: Document, items: list[dict], used: set[str]):
    for item in items:
        key = item.get("assetPath")
        if not key or key in used:
            continue
        used.add(key)
        caption = f"{item.get('_displayFigure', '图')} · {item.get('timeText', '')}".strip(" ·")
        add_picture(document, Path(key), caption=caption, width_cm=14.4, max_height_cm=8.6)


def add_table(document: Document, rows: list[list[str]]):
    rows = [row for row in rows if any(cell.strip() for cell in row)]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = clean_public_text(row[c_idx]) if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_text_run(p, text, bold=r_idx == 0)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.2)
            if r_idx == 0:
                set_cell_shading(cell, "F3F6FA")
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]):
    if not lines:
        return
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=6, after=8, line=1.15)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, east_asia="Consolas", latin="Consolas", size=9, color=(17, 24, 39))


def render_body(document: Document, markdown: str, evidence: list[dict]):
    lines = strip_evidence_map(markdown).splitlines()
    used_assets: set[str] = set()
    used_buckets: set[str] = set()
    paragraph_lines: list[str] = []
    bullet_lines: list[str] = []
    table_lines: list[str] = []
    code_lines: list[str] | None = None
    code_lang = ""
    skipped_h1 = False

    def flush_paragraph():
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = clean_public_text(" ".join(paragraph_lines))
        if text:
            p = document.add_paragraph()
            set_paragraph_spacing(p, after=7, line=1.55, first_indent=True)
            add_text_run(p, text)
        paragraph_lines = []

    def flush_bullets():
        nonlocal bullet_lines
        for item in bullet_lines:
            text = clean_public_text(item)
            if not text:
                continue
            p = document.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=4, line=1.45)
            add_text_run(p, text)
        bullet_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            rows = []
            for line in table_lines:
                if re.match(r"^\|\s*-+", line):
                    continue
                rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
            add_table(document, rows)
        table_lines = []

    def flush_code():
        nonlocal code_lines, code_lang
        if code_lines is not None and code_lang.lower() != "mermaid":
            add_code_block(document, code_lines)
        code_lines = None
        code_lang = ""

    def flush_all():
        flush_paragraph()
        flush_bullets()
        flush_table()

    for raw in lines:
        line = raw.rstrip()
        trimmed = line.strip()
        if re.search(r"证据文件|关键截图地图|结构化证据|逐帧抽取|证据图号", trimmed):
            flush_all()
            continue
        if code_lines is not None:
            if trimmed.startswith("```"):
                flush_code()
            else:
                code_lines.append(line)
            continue
        if trimmed.startswith("```"):
            flush_all()
            code_lang = trimmed.replace("```", "").strip()
            code_lines = []
            continue
        if not trimmed:
            flush_all()
            continue
        if re.match(r"^\|.*\|$", trimmed):
            flush_paragraph()
            flush_bullets()
            table_lines.append(trimmed)
            continue
        flush_table()
        heading = re.match(r"^(#{1,4})\s+(.+)$", trimmed)
        if heading:
            flush_all()
            level = min(3, len(heading.group(1)))
            text = clean_public_text(heading.group(2))
            if level == 1 and not skipped_h1:
                skipped_h1 = True
                continue
            if re.search(r"操作流程图|流程图|flowchart", text, re.I):
                continue
            add_heading(document, text, level)
            bucket = bucket_for_heading(text)
            if bucket and bucket not in used_buckets:
                used_buckets.add(bucket)
                add_evidence_block(document, pick_evidence_for_bucket(evidence, bucket), used_assets)
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", trimmed)
        if bullet:
            flush_paragraph()
            bullet_lines.append(bullet.group(1))
            continue
        numbered = re.match(r"^\d+[.、]\s+(.+)$", trimmed)
        if numbered:
            flush_paragraph()
            bullet_lines.append(numbered.group(1))
            continue
        paragraph_lines.append(trimmed)
    flush_code()
    flush_all()


def add_cover(document: Document, title: str, cover: Path | None):
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title_p, before=6, after=14, line=1.08)
    run = title_p.add_run(title)
    size = 34 if len(title) <= 22 else 28 if len(title) <= 38 else 23
    set_run_font(run, east_asia="宋体", latin="Times New Roman", size=size, bold=True, color=(32, 36, 43))
    dek = document.add_paragraph()
    set_paragraph_spacing(dek, after=18, line=1.6)
    add_text_run(dek, "围绕核心概念、操作步骤、关键设置与可复用清单整理，适合复习、修改和二次发布。")
    if cover:
        add_picture(document, cover, caption=None, width_cm=15.4, max_height_cm=13.2)
    document.add_page_break()


def add_mindmap(document: Document):
    image = SESSION_DIR / "learning_mindmap.png"
    if not image.exists():
        return
    add_heading(document, "精简思维导图", 1)
    add_picture(document, image, caption=None, width_cm=15.4, max_height_cm=20)
    document.add_page_break()


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def validate_docx(docx_path: Path):
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    bad = []
    if "\ufffd" in xml or "锟" in xml:
        bad.append("replacement characters")
    if re.search(r"\?{6,}", xml):
        bad.append("long question-mark run")
    if bad:
        raise RuntimeError("DOCX XML appears garbled: " + ", ".join(bad))


def main():
    summary_path = SESSION_DIR / "learning_summary.md"
    evidence_path = SESSION_DIR / "learning_evidence.json"
    if not summary_path.exists():
        raise FileNotFoundError(summary_path)
    if not evidence_path.exists():
        raise FileNotFoundError(evidence_path)

    markdown = read_text(summary_path)
    evidence_json = read_json(evidence_path)
    context = read_json(SESSION_DIR / "context.json")
    metadata = read_json(SESSION_DIR / "metadata.json")
    evidence = sorted_evidence(evidence_json)
    title = load_title(markdown, context, metadata)
    cover = resolve_cover(context, metadata)

    document = Document()
    configure_document(document)
    add_cover(document, title, cover)
    add_mindmap(document)
    render_body(document, markdown, evidence)

    out_path = SESSION_DIR / "learning_article.docx"
    document.save(out_path)
    validate_docx(out_path)
    print(json.dumps({"ok": True, "outPath": str(out_path), "title": title}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False, indent=2), file=sys.stderr)
        sys.exit(1)
