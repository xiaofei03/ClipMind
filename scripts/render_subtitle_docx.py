import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


if len(sys.argv) < 2:
    print("Usage: python render_subtitle_docx.py <session-dir>", file=sys.stderr)
    sys.exit(2)

SESSION_DIR = Path(sys.argv[1])


def read_text(path: Path) -> str:
    return path.read_text("utf-8", errors="replace") if path.exists() else ""


def read_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text("utf-8"))
    except Exception:
        return {}


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def add_run(paragraph, text: str, bold: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return run


def configure(document: Document):
    section = document.sections[0]
    section.top_margin = Pt(64)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def add_title(document: Document, title: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = add_run(p, title, True)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(28, 32, 39)
    p.paragraph_format.space_after = Pt(18)
    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lead, "字幕总结与逻辑树笔记")
    lead.paragraph_format.space_after = Pt(18)


def add_markdown(document: Document, markdown: str):
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            level = min(3, len(line) - len(line.lstrip("#")))
            text = clean_inline(line.lstrip("#").strip())
            if level == 1:
                continue
            heading = document.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            continue
        if re.match(r"^[-*+]\s+", line):
            p = document.add_paragraph(style="List Bullet")
            add_run(p, clean_inline(re.sub(r"^[-*+]\s+", "", line)))
            continue
        if re.match(r"^\d+\.\s+", line):
            p = document.add_paragraph(style="List Number")
            add_run(p, clean_inline(re.sub(r"^\d+\.\s+", "", line)))
            continue
        p = document.add_paragraph()
        add_run(p, clean_inline(line))
        p.paragraph_format.line_spacing = 1.45
        p.paragraph_format.space_after = Pt(6)


def validate(docx_path: Path):
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    if "\ufffd" in xml or "锟" in xml or re.search(r"\?{6,}", xml):
        raise RuntimeError("DOCX XML appears garbled")


def main():
    summary_path = SESSION_DIR / "subtitle_tree_summary.md"
    markdown = read_text(summary_path)
    if not markdown:
        raise FileNotFoundError(summary_path)
    context = read_json(SESSION_DIR / "context.json")
    title = context.get("metadata", {}).get("title") or "字幕总结"
    document = Document()
    configure(document)
    add_title(document, title)
    add_markdown(document, markdown)
    out_path = SESSION_DIR / "subtitle_tree_summary.docx"
    document.save(out_path)
    validate(out_path)
    print(json.dumps({"ok": True, "outPath": str(out_path)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False, indent=2), file=sys.stderr)
        sys.exit(1)
